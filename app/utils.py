import re
import math
import os
import shutil
import skimage.metrics as metrics
from skimage.io import imread
from skimage.color import rgb2gray
from urllib.parse import urlparse, parse_qs
from datetime import datetime
from docx import Document
from docx.shared import Inches
from cv2 import imread, cvtColor, COLOR_BGR2GRAY
from skimage.metrics import structural_similarity as ssim
from transliterate import translit
from gpt import send_request_gpt



def transliterate_russian(text):
    transliterated_text = translit(text, 'ru', reversed=True)
    return transliterated_text


def extract_link_from_message(text):
    url_pattern = re.compile(r'https?://\S+')

    match = re.search(url_pattern, text)

    if match:
        link = match.group(0)
        return link

    return None


def get_video_id(video_link):
    parsed_url = urlparse(video_link)
    query_params = parse_qs(parsed_url.query)
    video_id = query_params.get('v', [''])[0]
    return video_id


def compare_images(img1, img2):
    try:
        image1 = imread(img1)
        image2 = imread(img2)

        gray1 = cvtColor(image1, COLOR_BGR2GRAY)
        gray2 = cvtColor(image2, COLOR_BGR2GRAY)

        ssim_score, _ = ssim(gray1, gray2, full=True)

        return ssim_score

    except Exception as e:
        return 0


def compare_and_delete_images_in_folder(folder_path):
    image_files = os.listdir(folder_path)
    for i in range(len(image_files) - 1):
        for j in range(i + 1, len(image_files)):
            image1 = os.path.join(folder_path, image_files[i])
            image2 = os.path.join(folder_path, image_files[j])

            ssim_score = compare_images(image1, image2)

            if ssim_score > 0.85:
                # Delete one of the images (choose which one to delete)
                os.remove(image2)
                print(f"Deleted: {image2}")


def rename_files_in_folder(folder_path):
    filenames = os.listdir(folder_path)

    for i in range(len(filenames) - 1):
        current_file = filenames[i]
        next_file = filenames[i+1]

        current_timestamp = current_file.split('_')[2][:-4]
        next_timestamp = next_file.split('_')[2][:-4]

        if i != len(filenames)-1:
            new_filename = f"{current_file[:-4]}_{current_timestamp}-{next_timestamp}.jpg"
        else:
            new_filename = current_file

        os.rename(os.path.join(folder_path, current_file), os.path.join(folder_path, new_filename))

def create_new_docx_file(file_path):
    document = Document()
    document.save(file_path)

def add_text_to_docx(file_path, text, image=None):
    document = Document(file_path)
    document.add_paragraph(text)
    if image:
        document.add_picture(image, width=Inches(5.46), height=Inches(3))
    document.save(file_path)


def combine_sentence(screenshots_path, text_path, video_title,
                     output_path='/', annotation_len = 600, article_len = 6000):
    try:
        new_text_path = os.path.splitext(text_path)[0] + '.txt'

        os.rename(text_path, new_text_path)
    except Exception as e:
        print(e)
        new_text_path = text_path[0:4] + '.txt'

    sentences = []

    with open(new_text_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        for i in range(0, len(lines), 3):
            try:
                timecode = lines[i].strip()
                start_time, end_time = re.findall(r'\d{2}:\d{2}:\d{2}', timecode)
                text = lines[i + 1].strip()
                sentences.append({'start_time': start_time, 'end_time': end_time, 'text': text})
            except Exception as e:
                print(e)

    docx_path = 'result/output.docx'
    create_new_docx_file(docx_path)
    add_text_to_docx('result/output.docx', video_title)

    for filename in os.listdir(screenshots_path):
        try:
            screen_start_time = int(filename.split("_")[-1].replace(".jpg", "").split("-")[0])
            screen_end_time = int(filename.split("_")[-1].replace(".jpg", "").split("-")[1])
            screen_start_time, screen_end_time = sec_to_hhmmss(screen_start_time), sec_to_hhmmss(screen_end_time)
            matching_sentences = [sentence for sentence in sentences
                                  if screen_start_time <= sentence['start_time'] < screen_end_time]
            combined_text = " ".join([sentence['text'] for sentence in matching_sentences])
            # start_time_sec, end_time_sec = hhmmss_to_sec(screen_start_time), hhmmss_to_sec(screen_end_time)
            
            # text_to_file = f'{start_time_sec} {end_time_sec}\n{combined_text}\n\n'

            # _, new_combined_text = send_request_gpt(combined_text, annotation_len, article_len)

            text_to_file = f'{screen_start_time} {screen_end_time}\n{combined_text}\n\n'
            # text_to_file = f'{screen_start_time} {screen_end_time}\n{new_combined_text}\n\n'
            image = f'{screenshots_path}/{filename}'
            
            add_text_to_docx('result/output.docx', text_to_file, image)

            with open(f'{output_path}/output.txt', 'a', encoding='utf-8') as file:
                file.write(text_to_file)
            
        except Exception as e:
            print(e)


def sec_to_hhmmss(seconds):
    m, s = divmod(seconds, 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def hhmmss_to_sec(time):
    time_delta = datetime.timedelta(hours=int(time[:2]), minutes=int(time[3:5]), seconds=int(time[6:]))
    return time_delta.total_seconds()

def clear_directory(directory):
    for item in os.listdir(directory):
        item_path = os.path.join(directory, item)
        if os.path.isfile(item_path):
            os.remove(item_path)
        elif os.path.isdir(item_path):
            shutil.rmtree(item_path)

if __name__ == '__main__':
    folder_path = r"C:\Users\Tamara\Desktop\PROFBUH-hackathon-case-\app\screenshots\Vozvrat_tovarov_neplatel'schikom_NDS_v_1S_8.3_Buhgalt"
    txt_path = 'result/test.txt'
    output = 'result/'
    combine_sentence(folder_path, txt_path, output)
